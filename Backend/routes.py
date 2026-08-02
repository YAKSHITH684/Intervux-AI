from fastapi import APIRouter, UploadFile, File, Form
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from typing import List, Optional
import random
import os
import json
import requests
import re

from database.database import SessionLocal
from database.models import User

router = APIRouter()

GROQ_API_KEY = os.getenv("GROQ_API_KEY")
GROQ_URL = "https://api.groq.com/openai/v1/chat/completions"
GROQ_MODEL = "llama-3.3-70b-versatile"


def ask_groq(system: str, messages: list, max_tokens: int = 300) -> str:
    """Non-streaming call — still used by resume analysis, practice feedback, and the assistant chatbot."""
    if not GROQ_API_KEY:
        print("GROQ: No API key found in environment")
        return None
    try:
        payload = {
            "model": GROQ_MODEL,
            "max_tokens": max_tokens,
            "messages": [{"role": "system", "content": system}] + messages
        }
        res = requests.post(
            GROQ_URL,
            headers={"Authorization": f"Bearer {GROQ_API_KEY}", "Content-Type": "application/json"},
            json=payload,
            timeout=15
        )
        print(f"GROQ STATUS: {res.status_code} BODY: {res.text[:200]}")
        data = res.json()
        return data["choices"][0]["message"]["content"]
    except Exception as e:
        print(f"GROQ EXCEPTION: {e}")
        return None


def stream_groq(system: str, messages: list, max_tokens: int = 300):
    """Streams text chunks as Groq generates them — used by /chat so the first words
    reach the browser in well under a second instead of waiting for the full reply."""
    if not GROQ_API_KEY:
        print("GROQ: No API key found in environment")
        yield None
        return
    payload = {
        "model": GROQ_MODEL,
        "max_tokens": max_tokens,
        "messages": [{"role": "system", "content": system}] + messages,
        "stream": True
    }
    try:
        with requests.post(
            GROQ_URL,
            headers={"Authorization": f"Bearer {GROQ_API_KEY}", "Content-Type": "application/json"},
            json=payload,
            stream=True,
            timeout=15
        ) as res:
            got_any = False
            for line in res.iter_lines():
                if not line:
                    continue
                line = line.decode("utf-8")
                if not line.startswith("data: "):
                    continue
                data_str = line[len("data: "):]
                if data_str.strip() == "[DONE]":
                    break
                try:
                    chunk = json.loads(data_str)
                    delta = chunk["choices"][0]["delta"].get("content")
                    if delta:
                        got_any = True
                        yield delta
                except Exception:
                    continue
            if not got_any:
                yield None
    except Exception as e:
        print(f"GROQ STREAM EXCEPTION: {e}")
        yield None


# ─────────────────────────────────────────
# WARM-UP PING — called by every frontend page on load to wake
# a sleeping Render free-tier instance before the user's first
# real request (login, chat, resume upload, etc.)
# ─────────────────────────────────────────
@router.get("/ping")
def ping():
    return {"status": "awake"}


# ─────────────────────────────────────────
# AUTH
# ─────────────────────────────────────────
class RegisterRequest(BaseModel):
    email: str
    password: str

class LoginRequest(BaseModel):
    email: str
    password: str

class ForgotPasswordRequest(BaseModel):
    email: str
    new_password: str


@router.post("/register")
def register(data: RegisterRequest):
    db = SessionLocal()
    exists = db.query(User).filter(User.email == data.email).first()
    if exists:
        db.close()
        return {"success": False, "message": "User already exists"}
    user = User(email=data.email, password=data.password)
    db.add(user)
    db.commit()
    db.close()
    return {"success": True, "message": "Registration Successful"}


@router.post("/login")
def login(data: LoginRequest):
    db = SessionLocal()
    user = db.query(User).filter(
        User.email == data.email,
        User.password == data.password
    ).first()
    db.close()
    return {"success": bool(user)}


@router.post("/forgot-password")
def forgot_password(data: ForgotPasswordRequest):
    db = SessionLocal()
    user = db.query(User).filter(User.email == data.email).first()
    if not user:
        db.close()
        return {"success": False, "message": "Email not found"}
    user.password = data.new_password
    db.commit()
    db.close()
    return {"success": True, "message": "Password Updated Successfully"}


# ─────────────────────────────────────────
# RESUME ANALYSIS
# ─────────────────────────────────────────
@router.post("/analyze-resume")
async def analyze_resume(resume: UploadFile = File(...)):
    import pdfplumber
    import docx as python_docx

    content = await resume.read()
    filename = resume.filename.lower()
    text = ""

    if filename.endswith(".pdf"):
        with open("temp.pdf", "wb") as f:
            f.write(content)
        with pdfplumber.open("temp.pdf") as pdf:
            for page in pdf.pages:
                text += page.extract_text() or ""
    elif filename.endswith(".docx"):
        with open("temp.docx", "wb") as f:
            f.write(content)
        doc = python_docx.Document("temp.docx")
        for para in doc.paragraphs:
            text += para.text + " "
    else:
        try:
            text = content.decode("utf-8", errors="ignore")
        except Exception:
            text = ""

    if not text.strip():
        return {
            "ats_score": 0, "resume_score": 0, "readiness": 0, "skills": 0,
            "ai_skills_found": [],
            "ai_verdict": "Empty Resume (No content detected)",
            "ai_feedback": "No text could be extracted. Please upload a valid resume.",
            "summary": "No content detected.",
            "strengths": [],
            "suggestions": ["Upload a valid PDF/DOCX", "Make sure file is not blank"],
            "skill_breakdown": {"Technical Skills": 0, "Communication": 0, "Experience": 0, "Education": 0, "Projects": 0}
        }

    # Try Groq AI for full analysis
    import json as _json
    groq_system = """You are a professional ATS resume analyst. Analyze the resume and respond ONLY with valid JSON, no markdown, no explanation.
Schema:
{
  "ats_score": integer 0-100,
  "resume_score": integer 0-100,
  "readiness": integer 0-100,
  "skills": integer,
  "ai_skills_found": ["skill1", "skill2"],
  "ai_verdict": "Strong Resume / Moderate Resume / Weak Resume",
  "ai_feedback": "feedback string",
  "summary": "2-3 sentence summary",
  "strengths": ["s1","s2","s3"],
  "suggestions": ["s1","s2","s3","s4"],
  "skill_breakdown": {
    "Technical Skills": integer,
    "Communication": integer,
    "Experience": integer,
    "Education": integer,
    "Projects": integer
  }
}"""

    ai_reply = ask_groq(groq_system, [{"role": "user", "content": f"Analyze this resume:\n\n{text[:4000]}"}], max_tokens=1000)
    if ai_reply:
        try:
            clean = ai_reply.replace("```json", "").replace("```", "").strip()
            result = _json.loads(clean)
            return result
        except Exception as e:
            print(f"Groq JSON parse error: {e}")

    # Fallback keyword-based analysis
    text_lower = text.lower()
    ats = 30
    skill_map = {
        "python": "Python", "machine learning": "Machine Learning",
        "deep learning": "Deep Learning", "sql": "SQL",
        "fastapi": "FastAPI", "docker": "Docker", "aws": "AWS",
        "django": "Django", "flask": "Flask", "javascript": "JavaScript",
        "react": "React", "node": "Node.js", "java": "Java",
        "c++": "C++", "tensorflow": "TensorFlow", "pytorch": "PyTorch",
        "git": "Git", "linux": "Linux", "mongodb": "MongoDB",
        "rest api": "REST API", "html": "HTML", "css": "CSS"
    }
    skills_found = []
    for key, val in skill_map.items():
        if key in text_lower:
            skills_found.append(val)
            ats = min(ats + 4, 100)

    if "project" in text_lower: ats = min(ats + 5, 100)
    if "experience" in text_lower: ats = min(ats + 5, 100)
    if "github" in text_lower: ats = min(ats + 3, 100)
    if "internship" in text_lower: ats = min(ats + 5, 100)

    feedback = []
    if "project" not in text_lower: feedback.append("Add real-world projects")
    if "github" not in text_lower: feedback.append("Include your GitHub profile")
    if "experience" not in text_lower: feedback.append("Mention internships or work experience")
    if len(skills_found) < 3: feedback.append("Add more technical skills")
    if not feedback: feedback.append("Well-structured resume")

    verdict = "Strong Resume (Great skill set)" if len(skills_found) >= 6 else \
              "Moderate Resume (Needs improvement)" if len(skills_found) >= 3 else \
              "Weak Resume (Add more skills)"

    return {
        "ats_score": ats,
        "resume_score": max(ats - 5, 0),
        "readiness": max(ats - 10, 0),
        "skills": len(skills_found),
        "ai_skills_found": skills_found,
        "ai_verdict": verdict,
        "ai_feedback": " | ".join(feedback),
        "summary": f"Candidate has {len(skills_found)} detected skills. ATS compatibility is {ats}%.",
        "strengths": skills_found[:3] if skills_found else ["Skills not detected"],
        "suggestions": feedback[:4],
        "skill_breakdown": {
            "Technical Skills": min(30 + len(skills_found) * 5, 100),
            "Communication": 65,
            "Experience": 60 if "experience" in text_lower else 40,
            "Education": 70 if "education" in text_lower or "degree" in text_lower else 50,
            "Projects": 75 if "project" in text_lower else 35
        }
    }


# ─────────────────────────────────────────
# AI INTERVIEW CHAT — Groq AI + fallback
# ─────────────────────────────────────────
class ChatRequest(BaseModel):
    message: str
    history: List[dict] = []
    mode: str = "technical"
    difficulty: str = "medium"  # easy | medium | hard

SYSTEM_PROMPTS = {
    "technical": """You are a senior software engineer conducting a technical interview.
Ask one focused technical question at a time. After the candidate answers:
1. Give specific feedback on their answer (2-3 sentences)
2. Adjust difficulty based on their response quality
3. Ask the next question

If the answer is weak or too short, ask a simpler follow-up or hint.
If the answer is strong, increase difficulty slightly.
Keep responses concise. Do not repeat questions already asked.""",

    "hr": """You are an experienced HR manager conducting a behavioral interview.
Ask one behavioral question at a time using STAR format.
After each answer:
1. Give encouraging, specific feedback (2-3 sentences)
2. Note what was good and what could be improved
3. Ask the next question

Adapt questions based on the candidate's experience level shown in their answers.""",

    "system": """You are a principal engineer conducting a system design interview.
Ask one system design question at a time.
After each answer:
1. Give technical feedback highlighting strengths and gaps
2. Ask a follow-up to dig deeper OR move to next topic
3. Guide with hints if the candidate is stuck

Focus on scalability, trade-offs, and real-world constraints.""",

    "dsa": """You are a FAANG engineer conducting a DSA coding interview.
Present one problem at a time.
After each answer:
1. Evaluate correctness, time complexity, space complexity
2. Give specific feedback
3. If wrong, give a hint rather than the answer
4. Move to next problem when ready

Start easy, increase difficulty based on performance."""
}

FALLBACK_QUESTIONS = {
    "technical": [
        "Tell me about yourself and your technical background.",
        "What programming languages are you most comfortable with?",
        "Explain the difference between a list and a tuple in Python.",
        "What is OOP and what are its main principles?",
        "Explain what a REST API is and how it works.",
        "What is the difference between SQL and NoSQL databases?",
        "How does Git version control work?",
        "What is time complexity and why does it matter?",
        "Describe a challenging technical problem you solved.",
        "Where do you see yourself in 5 years technically?"
    ],
    "hr": [
        "Tell me about yourself.",
        "What are your greatest strengths?",
        "What is your biggest weakness and how do you overcome it?",
        "Describe a time you worked in a team and faced a conflict.",
        "Tell me about a project you are most proud of.",
        "How do you handle pressure and tight deadlines?",
        "Why do you want to work at our company?",
        "Where do you see yourself in 5 years?",
        "What motivates you to do your best work?",
        "Why should we hire you over other candidates?"
    ],
    "system": [
        "How would you design a URL shortener like bit.ly?",
        "Design a scalable chat application like WhatsApp.",
        "How would you design a content delivery network?",
        "Explain how you would design a ride-sharing system like Uber.",
        "How would you design a recommendation system?",
        "What are trade-offs between SQL and NoSQL for a social media app?",
        "How would you ensure high availability in a distributed system?",
        "Explain load balancing and when you would use it.",
        "How would you design a notification system for millions of users?",
        "What is the CAP theorem and how does it affect system design?"
    ],
    "dsa": [
        "Reverse a string without using built-in functions.",
        "Find the largest element in an array.",
        "Check if a string is a palindrome.",
        "Find duplicates in an array.",
        "Implement a binary search algorithm.",
        "Explain how a linked list works and implement insertion.",
        "What is a stack and implement push/pop operations.",
        "Find the Fibonacci number at position N using dynamic programming.",
        "Explain BFS and DFS graph traversal.",
        "Find two numbers in an array that add up to a target sum."
    ]
}

FALLBACK_FEEDBACK = [
    "Good answer! You demonstrated clear thinking.",
    "Nice response! Try to include a specific example next time.",
    "Well explained! Adding metrics or numbers makes it stronger.",
    "Good structure! Keep answers concise and to the point.",
    "Great effort! Consider using the STAR method for behavioral questions.",
    "Solid answer! A real-world example would make it even better.",
    "Good thinking! Practice elaborating on your technical decisions.",
    "Nice work! Confidence in delivery is key in real interviews."
]


@router.post("/chat")
def chat(data: ChatRequest):
    """Streams the reply back as plain text chunks. The frontend is responsible for
    appending each chunk to `history` once the stream finishes (see interview.html)."""
    mode = data.mode if data.mode in SYSTEM_PROMPTS else "technical"
    history = list(data.history)
    msg = data.message.strip()

    # First message — start interview
    if not history:
        mode_label = {
            "technical": "Technical", "hr": "HR / Behavioral",
            "system": "System Design", "dsa": "DSA"
        }.get(mode, "Technical")

        prompt_messages = [{
            "role": "user",
            "content": f"Start the interview. Mode: {mode_label}. Greet me warmly and ask the first question."
        }]

        def gen():
            got_any = False
            for chunk in stream_groq(SYSTEM_PROMPTS[mode], prompt_messages):
                if chunk is None and not got_any:
                    questions = FALLBACK_QUESTIONS[mode]
                    yield (
                        f"Welcome to the AI Interview Simulator!\n\n"
                        f"Mode: {mode_label}\n\n"
                        f"I will ask you 10 questions. Take your time and answer clearly.\n\n"
                        f"Tip: Be specific, use examples, and stay confident!\n\n"
                        f"---\n"
                        f"Question 1 of 10:\n\n{questions[0]}"
                    )
                    return
                if chunk:
                    got_any = True
                    yield chunk

        return StreamingResponse(gen(), media_type="text/plain")

    # Count answered questions
    user_turns = [h for h in history if h.get("role") == "user"]
    answered_count = len(user_turns)

    # Add user message
    history.append({"role": "user", "content": msg})

    # Interview complete after 10 questions
    if answered_count >= 10:
        def gen_done():
            yield (
                "Interview Complete!\n\n"
                "You answered all 10 questions.\n\n"
                "Great job! Keep practicing to improve your confidence.\n\n"
                "Click New to start another round."
            )
        return StreamingResponse(gen_done(), media_type="text/plain")

    def gen():
        got_any = False
        for chunk in stream_groq(SYSTEM_PROMPTS[mode], history[-6:]):  # last 6 turns for context
            if chunk is None and not got_any:
                questions = FALLBACK_QUESTIONS[mode]
                fb = random.choice(FALLBACK_FEEDBACK)
                next_q_num = answered_count + 1
                next_question = questions[min(answered_count, len(questions) - 1)]
                yield (
                    f"Feedback: {fb}\n\n"
                    f"---\n"
                    f"Question {next_q_num} of 10:\n\n{next_question}"
                )
                return
            if chunk:
                got_any = True
                yield chunk

    return StreamingResponse(gen(), media_type="text/plain")


# ─────────────────────────────────────────
# PRACTICE QUESTIONS
# ─────────────────────────────────────────
class PracticeRequest(BaseModel):
    topic: str
    count: int = 3

class FeedbackRequest(BaseModel):
    question: str
    answer: str
    topic: str

PRACTICE_QUESTIONS = {
    "arrays": [
        {"question": "Find the maximum subarray sum (Kadane's Algorithm).", "hint": "Keep a running sum, reset when negative.", "difficulty": "Medium"},
        {"question": "Remove duplicates from a sorted array in-place.", "hint": "Use two pointers.", "difficulty": "Easy"},
        {"question": "Rotate an array to the right by K steps.", "hint": "Reverse the whole array, then parts.", "difficulty": "Medium"},
    ],
    "python": [
        {"question": "What is the difference between a list and a tuple?", "hint": "Think mutability.", "difficulty": "Easy"},
        {"question": "Explain decorators in Python with an example.", "hint": "Functions that wrap other functions.", "difficulty": "Medium"},
        {"question": "What are Python generators and when would you use them?", "hint": "Think memory efficiency and yield.", "difficulty": "Medium"},
    ],
    "sql": [
        {"question": "Write a query to find the second highest salary.", "hint": "Use LIMIT/OFFSET or subquery.", "difficulty": "Medium"},
        {"question": "Explain the difference between INNER JOIN and LEFT JOIN.", "hint": "Think about what rows are included.", "difficulty": "Easy"},
        {"question": "What is database indexing and why is it important?", "hint": "Think about query speed.", "difficulty": "Easy"},
    ],
    "default": [
        {"question": "Explain the concept of recursion with an example.", "hint": "Think base case and recursive case.", "difficulty": "Easy"},
        {"question": "What is the difference between stack and queue?", "hint": "Think LIFO vs FIFO.", "difficulty": "Easy"},
        {"question": "Explain time complexity with an example.", "hint": "Think Big O notation.", "difficulty": "Medium"},
    ]
}

@router.post("/practice/questions")
def generate_questions(data: PracticeRequest):
    topic_lower = data.topic.lower()
    if "array" in topic_lower or "string" in topic_lower:
        questions = PRACTICE_QUESTIONS["arrays"]
    elif "python" in topic_lower or "oop" in topic_lower:
        questions = PRACTICE_QUESTIONS["python"]
    elif "sql" in topic_lower or "database" in topic_lower:
        questions = PRACTICE_QUESTIONS["sql"]
    else:
        questions = PRACTICE_QUESTIONS["default"]
    return {"success": True, "questions": questions[:data.count]}


@router.post("/practice/feedback")
def get_feedback(data: FeedbackRequest):
    # Try Groq AI for personalized feedback
    ai_reply = ask_groq(
        "You are an expert interviewer giving feedback on an interview answer. Be specific, constructive and encouraging. Keep it to 3-4 sentences.",
        [{"role": "user", "content": f"Topic: {data.topic}\nQuestion: {data.question}\nAnswer: {data.answer}\n\nGive feedback."}],
        max_tokens=300
    )
    if ai_reply:
        return {"success": True, "feedback": ai_reply}

    # Fallback
    ans = data.answer.strip()
    if len(ans) < 10:
        feedback = "Your answer is too short. Please elaborate with examples."
    elif len(ans) < 50:
        feedback = "Decent start! Try to expand with a concrete example or code snippet."
    elif len(ans) < 150:
        feedback = "Good answer! You covered the basics. Adding real-world examples would make it stronger."
    else:
        feedback = "Excellent detailed answer! You demonstrated strong understanding. Keep it up!"
    return {"success": True, "feedback": feedback}


# ─────────────────────────────────────────
# DASHBOARD ASSISTANT CHATBOT
# ─────────────────────────────────────────
class AssistantRequest(BaseModel):
    message: str
    history: List[dict] = []

ASSISTANT_SYSTEM = """You are a friendly AI assistant and interview coach. 
- Answer any question naturally and conversationally
- For casual messages like "how are you", "hi", "thanks" — respond naturally and warmly
- For interview-related questions — give helpful, practical advice
- Keep all responses short and simple (2-3 sentences max)
- Be friendly, warm and human-like"""

ASSISTANT_KB = {
    "tell me about yourself": "Structure your answer in 3 parts: 1) Who you are, 2) Your key skills/experience, 3) Why you are excited about this role. Keep it under 2 minutes. Example: I am a computer science graduate with strong Python and ML skills. I have built 3 projects including an AI chatbot. I am passionate about solving real problems with AI.",
    "ats": "ATS (Applicant Tracking System) score measures how well your resume matches a job description. Tips to improve: use keywords from the job posting, use standard section headings like Experience and Education, avoid tables and graphics, and save as PDF.",
    "technical interview": "To prepare for technical interviews: 1) Practice DSA on LeetCode daily, 2) Review core CS concepts (OOP, OS, Networks, DB), 3) Build projects to discuss, 4) Practice explaining your thought process out loud, 5) Do mock interviews.",
    "weakness": "When answering What is your weakness: Choose a real weakness that is not critical to the job. Show self-awareness and explain steps you are taking to improve. Example: I sometimes spend too much time perfecting code. I have learned to set time limits and prioritize delivery.",
    "strength": "For strengths: Pick 2-3 relevant to the role. Back each with a specific example. Example: My strength is problem-solving. In my last project I debugged a critical API issue that was blocking the team by using systematic logging.",
    "salary": "For salary negotiation: Research market rates on Glassdoor and LinkedIn first. Give a range based on your research. Say: Based on my research and experience, I am looking for this range. I am open to discussion based on the full package.",
    "resume": "Resume tips: Keep it to 1 page, use action verbs (built, designed, improved), add metrics (increased performance by 30%), include GitHub link, list relevant skills, and tailor it for each job.",
    "hr interview": "For HR interviews use the STAR method: Situation, Task, Action, Result. Be honest, show enthusiasm, research the company beforehand, and prepare 3-5 stories that demonstrate your skills.",
    "project": "When explaining your project: 1) What problem it solves, 2) Technologies used, 3) Your specific role and contribution, 4) Challenges faced and how you solved them, 5) Results and impact.",
    "pressure": "To answer How do you handle pressure: Give a real example. In my final year project, we had a 48-hour deadline. I broke the task into smaller chunks, prioritized critical features, and communicated progress to my team. We delivered on time.",
}

@router.post("/assistant")
def assistant(data: AssistantRequest):
    msg = data.message.lower().strip()
    history = data.history.copy()
    history.append({"role": "user", "content": data.message})

    # Try Groq AI first
    ai_reply = ask_groq(ASSISTANT_SYSTEM, history[-6:], max_tokens=400)
    if ai_reply:
        history.append({"role": "assistant", "content": ai_reply})
        return {"reply": ai_reply, "history": history}

    # Fallback knowledge base
    reply = None
    for key, answer in ASSISTANT_KB.items():
        if any(word in msg for word in key.split()):
            reply = answer
            break

    if not reply:
        if any(w in msg for w in ["hi", "hello", "hey"]):
            reply = "Hello! I am your AI Interview Assistant. Ask me about interview tips, resume advice, or practice questions!"
        elif any(w in msg for w in ["thank", "thanks"]):
            reply = "You are welcome! Best of luck with your interview. You have got this!"
        elif any(w in msg for w in ["interview", "prepare", "preparation"]):
            reply = "Interview Preparation Tips:\n\n1. Research the company thoroughly\n2. Practice common questions out loud\n3. Prepare 5-6 STAR stories\n4. Review your resume and projects\n5. Dress professionally and arrive early\n6. Ask thoughtful questions at the end"
        elif any(w in msg for w in ["nervous", "anxiety", "scared", "fear"]):
            reply = "Interview nerves are completely normal! Tips to calm down:\n\n1. Prepare thoroughly, confidence comes from preparation\n2. Take deep breaths before entering\n3. Remember they want you to succeed\n4. Treat it as a conversation, not an exam\n5. Practice mock interviews with friends"
        elif any(w in msg for w in ["dsa", "algorithm", "leetcode", "coding"]):
            reply = "DSA Interview Tips:\n\n1. Start with easy LeetCode problems\n2. Master arrays, strings, hashmaps first\n3. Learn binary search and two pointers\n4. Practice trees and graphs\n5. Always explain your approach before coding\n6. Discuss time and space complexity"
        elif any(w in msg for w in ["system design", "design"]):
            reply = "System Design Tips:\n\n1. Clarify requirements first\n2. Estimate scale (users, requests per second)\n3. Draw high-level architecture\n4. Discuss database choices (SQL vs NoSQL)\n5. Talk about caching and load balancing\n6. Address failure scenarios"
        else:
            reply = "I can help you with:\n\n- Tell me about yourself\n- Technical interview tips\n- Resume and ATS advice\n- HR and behavioral questions\n- Salary negotiation\n- DSA and System Design tips\n- Handling interview nerves\n\nWhat would you like to know?"

    history.append({"role": "assistant", "content": reply})
    return {"reply": reply, "history": history}


# ─────────────────────────────────────────
# AI ANALYSIS ENDPOINT
# ─────────────────────────────────────────
class AIAnalysisRequest(BaseModel):
    resume_text: str

@router.post("/ai-analysis")
def ai_analysis(data: AIAnalysisRequest):
    import json as _json
    system = """You are a professional resume analyst. Analyze the resume deeply.
Respond ONLY with valid JSON, no markdown, no explanation.
Schema:
{
  "ats_score": integer,
  "readiness_score": integer,
  "skills": ["skill1", "skill2"],
  "suggestions": ["s1","s2","s3","s4"],
  "summary": "string",
  "skill_breakdown": {
    "Technical Skills": integer,
    "Communication": integer,
    "Experience": integer,
    "Education": integer,
    "Projects": integer
  },
  "strengths": ["s1","s2","s3"],
  "gaps": ["g1","g2"]
}"""

    ai_reply = ask_groq(system, [{"role": "user", "content": data.resume_text[:4000]}], max_tokens=1000)
    if ai_reply:
        try:
            clean = ai_reply.replace("```json", "").replace("```", "").strip()
            result = _json.loads(clean)
            return {"success": True, **result}
        except Exception as e:
            return {"success": False, "error": f"Parse error: {e}"}
    return {"success": False, "error": "Groq unavailable"}


# ─────────────────────────────────────────
# JOB DESCRIPTION MATCHING
# ─────────────────────────────────────────
# Shared keyword bank used by the fallback matcher (kept broader than the
# analyze-resume skill_map since job descriptions mention tools/soft skills
# that resumes don't always phrase the same way).
MATCH_SKILLS_DB = [
    "python", "java", "c++", "c#", "javascript", "typescript", "go", "rust",
    "sql", "nosql", "mongodb", "postgresql", "mysql", "redis",
    "machine learning", "deep learning", "nlp", "computer vision",
    "tensorflow", "pytorch", "scikit-learn", "pandas", "numpy",
    "django", "flask", "fastapi", "spring", "node.js", "express",
    "react", "angular", "vue", "html", "css", "tailwind",
    "docker", "kubernetes", "aws", "azure", "gcp", "ci/cd", "jenkins",
    "git", "github", "linux", "rest api", "graphql", "microservices",
    "agile", "scrum", "communication", "leadership", "problem solving",
    "teamwork", "data structures", "algorithms", "system design",
    "testing", "unit testing", "debugging", "excel", "tableau", "power bi",
]


def _extract_keywords(text: str) -> set:
    text_lower = text.lower()
    return {kw for kw in MATCH_SKILLS_DB if kw in text_lower}


class JobMatchRequest(BaseModel):
    resume_text: str
    job_description: str


def _fallback_job_match(resume_text: str, job_description: str) -> dict:
    resume_kw = _extract_keywords(resume_text)
    jd_kw = _extract_keywords(job_description)

    matched = sorted(resume_kw & jd_kw)
    missing = sorted(jd_kw - resume_kw)

    match_score = round((len(matched) / len(jd_kw)) * 100) if jd_kw else 50

    # crude keyword-density proxy for the two other headline numbers
    resume_words = re.findall(r"[a-zA-Z]+", resume_text.lower())
    jd_words = re.findall(r"[a-zA-Z]+", job_description.lower())
    overlap_words = set(resume_words) & set(jd_words)
    ats_score = min(round(match_score * 0.8 + len(overlap_words) / 5), 100)

    suggestions = []
    if missing:
        suggestions.append(f"Add these missing keywords if genuinely applicable: {', '.join(missing[:6])}")
    if len(matched) < 3:
        suggestions.append("Mirror more of the job description's language in your skills/experience sections")
    if "project" not in resume_text.lower():
        suggestions.append("Add project details that relate directly to this role")
    if not suggestions:
        suggestions.append("Strong overlap with this role — tailor your summary line to reinforce it")

    verdict = (
        "Strong Match" if match_score >= 70 else
        "Moderate Match" if match_score >= 40 else
        "Weak Match"
    )

    return {
        "success": True,
        "match_score": match_score,
        "ats_score": ats_score,
        "verdict": verdict,
        "matched_skills": matched,
        "missing_skills": missing,
        "summary": f"{len(matched)} of {len(jd_kw) or 1} keywords from the job description were found in the resume.",
        "suggestions": suggestions[:4],
        "category_breakdown": {
            "Skills Match": match_score,
            "Keyword Coverage": ats_score,
            "Experience Relevance": 60 if "experience" in resume_text.lower() else 40,
            "Project Relevance": 65 if "project" in resume_text.lower() else 35,
        },
    }


@router.post("/match-job")
def match_job(data: JobMatchRequest):
    """Compares a resume against a job description and returns a match score,
    matched/missing keywords, and tailored suggestions. Tries Groq first,
    falls back to keyword-overlap scoring if the AI call fails."""
    import json as _json

    system = """You are an expert ATS and recruiting analyst. Compare the resume to the job description and respond ONLY with valid JSON, no markdown, no explanation.
Schema:
{
  "match_score": integer 0-100,
  "ats_score": integer 0-100,
  "verdict": "Strong Match / Moderate Match / Weak Match",
  "matched_skills": ["skill1", "skill2"],
  "missing_skills": ["skill1", "skill2"],
  "summary": "2-3 sentence summary of fit",
  "suggestions": ["s1","s2","s3","s4"],
  "category_breakdown": {
    "Skills Match": integer,
    "Keyword Coverage": integer,
    "Experience Relevance": integer,
    "Project Relevance": integer
  }
}"""

    user_content = (
        f"RESUME:\n{data.resume_text[:3500]}\n\n"
        f"JOB DESCRIPTION:\n{data.job_description[:2000]}"
    )

    ai_reply = ask_groq(system, [{"role": "user", "content": user_content}], max_tokens=1000)
    if ai_reply:
        try:
            clean = ai_reply.replace("```json", "").replace("```", "").strip()
            result = _json.loads(clean)
            return {"success": True, **result}
        except Exception as e:
            print(f"Groq JSON parse error (match-job): {e}")

    return _fallback_job_match(data.resume_text, data.job_description)


@router.post("/match-job-file")
async def match_job_file(
    resume: UploadFile = File(...),
    job_description: str = Form(...),
):
    """Same as /match-job but accepts a resume file (PDF/DOCX/TXT) instead of raw text."""
    import pdfplumber
    import docx as python_docx

    content = await resume.read()
    filename = resume.filename.lower()
    text = ""

    if filename.endswith(".pdf"):
        with open("temp_match.pdf", "wb") as f:
            f.write(content)
        with pdfplumber.open("temp_match.pdf") as pdf:
            for page in pdf.pages:
                text += page.extract_text() or ""
    elif filename.endswith(".docx"):
        with open("temp_match.docx", "wb") as f:
            f.write(content)
        doc = python_docx.Document("temp_match.docx")
        for para in doc.paragraphs:
            text += para.text + " "
    else:
        try:
            text = content.decode("utf-8", errors="ignore")
        except Exception:
            text = ""

    if not text.strip():
        return {"success": False, "error": "Could not extract text from resume file"}

    return match_job(JobMatchRequest(resume_text=text, job_description=job_description))